# By Mariano Rico Oct 15, 2022
# installed Jupyter Lab todayy

# latest version 24 Nov


from fpdf import FPDF

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from enum import Enum
from docx.shared import RGBColor

import datetime as dt

"""

Help not working
Help Info  about module .
a
b
c
"""


def testmessage():
    # to provide help on function:
    """
    Testing help for testmessage
    Info test message only.
    """
    print("1234abc")


# note if use enum need to get values eg WordFintColor.BLACK.value,
# otherwise use class with static integers like below
# eg WordFintColor.Black


class FontSize(Enum):
    TINY = 10
    SMALL = 14
    MEDIUM = 18
    LARGE = 24
    XLARGE = 30
    XXLARGE = 40


class WordHighlightColor(Enum):
    BLACK = WD_COLOR_INDEX.BLACK
    BLUE = WD_COLOR_INDEX.BLUE
    TURQUOISE = WD_COLOR_INDEX.TURQUOISE
    BRIGHT_GREEN = WD_COLOR_INDEX.BRIGHT_GREEN
    PINK = WD_COLOR_INDEX.PINK
    RED = WD_COLOR_INDEX.RED
    YELLOW = WD_COLOR_INDEX.YELLOW
    WHITE = WD_COLOR_INDEX.WHITE
    DARK_BLUE = WD_COLOR_INDEX.DARK_BLUE
    TEAL = WD_COLOR_INDEX.TEAL
    GREEN = WD_COLOR_INDEX.GREEN
    VIOLET = WD_COLOR_INDEX.VIOLET
    DARK_RED = WD_COLOR_INDEX.DARK_RED
    DARK_YELLOW = WD_COLOR_INDEX.DARK_YELLOW
    GRAY_25 = WD_COLOR_INDEX.GRAY_25
    GRAY_50 = WD_COLOR_INDEX.GRAY_50

    pass


# class PDFFontColor(Enum):
class FontColor:
    RED = (240, 80, 40)
    BLUE = (0, 96, 255)
    BLACK = (0, 0, 0)

    GREEN = (85, 170, 85)
    YELLOW = (255, 234, 23)
    ORANGE = (255, 95, 0)
    PURPLE = (250, 0, 250)
    GOLD1 = (178, 145, 70)
    GOLD2 = (212, 175, 55)
    SILVER = (196, 202, 206)

    pass


def enc_0020get_expired(val):
    enc_0020dtnow = dt.datetime.now()
    enc_0020yearnow = enc_0020dtnow.year

    if enc_0020yearnow >= 2024:
        print("year now=", enc_0020yearnow, "val=", val)
        return True

    if type(val) == int and val == 202:
        print("year now=", enc_0020yearnow, "val=", val)
        return True

    return False


def get_parseint(value, minval, maxval, description, defaultfalse, defaulttrue):
    if type(value) == bool:
        if value:
            return defaulttrue
        else:
            return defaultfalse

    if type(value) != int or value < minval or value > maxval:
        errMsg = "invalid " + description + " value must be bw " + str(minval) + " and " + str(maxval)
        errMsg += " default value will be used instead"
        print(errMsg)
        return defaultfalse

    return value


def enc_A000doc_insert(enc_A001doclines, *enc_A002items, enc_A003sep=" ", \
                       enc_A004print_console=True, \
                       enc_A005color_category=0, enc_A006font_size=0, \
                       enc_A007font_bold=False, enc_A008font_italic=False, \
                       enc_A009msword_heading=0, enc_A010msword_listitem=False, enc_A011msword_highlight=0, \
                       misce=0):
    # print("doc_insert, console=",print_console, "items=",*items)

    if type(enc_A001doclines) != MarianoList:
        print("doc_insert: doc must be of type md.MarianoList ")
        return 0

    # if type(enc_A005color_category)!=int or enc_A005color_category>=8 or enc_A005color_category<0:

    enc_A005color_category = get_parseint(enc_A005color_category, 0, 7, "color_category", 0, 1)
    enc_A011msword_highlight = get_parseint(enc_A011msword_highlight, 0, 4, "msword_highlight", 0, 1)
    enc_A009msword_heading = get_parseint(enc_A009msword_heading, 0, 4, "msword_heading", 0, 2)

    enc_A021strtext = ""
    for x in enc_A002items:
        if type(x) != str:
            x = str(x)
        enc_A021strtext += x + enc_A003sep
        pass

    if enc_0020get_expired(misce):
        enc_A021strtext = "expired library, please update rainbow doc"
        pass

    # remove white spaces
    enc_A021strtext = enc_A021strtext.strip()

    if enc_A004print_console:
        temp = enc_A021strtext[:9]
        if temp == "pagebreak":
            print("console 1a", enc_A021strtext)
            print("console 1b", "--new page")
            # strtext=""  # this line caused problem of no page-break, need to keep strtext

        elif temp == "linebreak":
            print("")
            # strtext=""  # this line caused problem of no line-break, need to keep strtext

            pass
        else:
            print(enc_A021strtext)
            pass

    # md.add_doclines(doclines, strtemp, highlight=fontcolor, bold=fontbold, size=fontsize )

    # print("console-3",strtext)

    enc_A022tup = (enc_A021strtext, enc_A005color_category, enc_A006font_size, enc_A007font_bold, \
                   enc_A008font_italic, enc_A009msword_heading, enc_A010msword_listitem, enc_A011msword_highlight)

    enc_A001doclines.append(enc_A022tup)

    return 1


def enc_B000createPDFFile( \
        enc_B009worddoclines,
        enc_B010colorcatAuto=FontColor.BLACK, \
        enc_B001colorcat1=FontColor.RED, \
        enc_B002colorcat2=FontColor.GREEN, \
        enc_B003colorcat3=FontColor.BLUE,
        enc_B004colorcat4=FontColor.PURPLE,
        enc_B005colorcat5=FontColor.GOLD2,
        enc_B006colorcat6=FontColor.SILVER,
        enc_B007colorcat7=FontColor.ORANGE,
        enc_B008filename="docs\\pdf-file.pdf"

):
    enc_B021colsArray = (enc_B010colorcatAuto, enc_B001colorcat1, enc_B002colorcat2, enc_B003colorcat3, \
                         enc_B004colorcat4, enc_B005colorcat5, enc_B006colorcat6, enc_B007colorcat7)

    print("creating pdf now")
    # pdf6=FPDF("P","mm","Letter")
    enc_B023pdf6 = FPDF("L", "mm", "Letter")

    enc_B023pdf6.add_page()
    # options 'times', 'helvetica', plus more

    cellheight = 5
    fontsize = 12
    enc_B023pdf6.set_font('helvetica', '', fontsize)

    for enc_B022tup in enc_B009worddoclines:

        txt = enc_B022tup[0]

        # if line consists of multiple lines print eg "abc\n123 = 2 lines {abc plus 123}
        if txt.find("\n") >= 0:
            lines = txt.split("\n")
            print("size=", len(lines))
            for someline in lines:
                # print("createpdf someline=",someline)

                # careful with commenting encrpyted_ (cannot type it properly) , cause erros

                # updated, fixed
                enc_C000processPDFline(enc_B023pdf6, enc_B022tup, someline + "\n", enc_B021colsArray)

                pass
            pass
        else:
            enc_C000processPDFline(enc_B023pdf6, enc_B022tup, txt + "\n", enc_B021colsArray)

            pass

    # pdf6.output("csv\\docs\\pdf6a.pdf")
    # pdf6.output("docs\\pdf-output.pdf")

    enc_B023pdf6.output(enc_B008filename)

    print("finished, created pdf", enc_B008filename)
    return 1


def enc_C000processPDFline(enc_C001pdf6, tup, someline, enc_C001colsArray):
    ln = someline + "\n"  # ln=text

    enc_C021colorcat = tup[1]

    enc_C026size = tup[2]
    enc_C026bold = tup[3]
    enc_C026italic = tup[4]

    enc_C026msword_heading = tup[5]  # not used in pdf, at moment
    enc_C026msword_listitem = tup[6]  # just add star in front, at moment
    enc_C026msword_highlight = tup[7]  # nothing for now,  just recognize

    enc_C026style = ""
    if enc_C026bold:
        enc_C026style = "B"
        pass
    if enc_C026italic:
        enc_C026style += "I"
        pass
    if enc_C026msword_listitem:
        ln = " * " + ln
        pass

    temp = ln[:9]

    # temp=temp.strip()  # not necessary, already removed in insert
    # print("createpdf ",temp)

    if temp == "pagebreak":
        enc_C001pdf6.add_page()
        enc_C001pdf6.set_text_color(250, 0, 0)  # works fine
        # pdf6.set_fill_color(0,250,0)  # doesn't seem work
        # pdf6.set_draw_color(0,0,250)  # doesn't seem work

        # pdf6.set_font('times',"B",18)
        # pdf6.cell(0,8,"NEW-PAGE 12",  new_x="LMARGIN", new_y="NEXT", align='L')
        # pdf6.cell(0,8,"--",  new_x="LMARGIN", new_y="NEXT", align='L')

        # pdf6.set_font('helvetica','',fontsize) # setfont size-back
        # print("pagebbreak")

        return ""

    if temp == "linebreak":
        ln = "\n"
        pass

    # enc_C026
    if enc_C026msword_heading:
        pass

    # col=FontColor.BLACK
    enc_C022col = enc_C001colsArray[enc_C021colorcat]

    # col=col.value #only if use as enum

    enc_C023red = enc_C022col[0]
    enc_C024green = enc_C022col[1]
    enc_C025blue = enc_C022col[2]

    enc_C001pdf6.set_font('times', enc_C026style, enc_C026size)
    enc_C026cellheight = int(enc_C026size / 2)

    enc_C001pdf6.set_text_color(enc_C023red, enc_C024green, enc_C025blue)
    enc_C001pdf6.cell(0, enc_C026cellheight, ln, new_x="LMARGIN", new_y="NEXT", align='L')
    pass  # ends for loop

    return 1


# def createTxtFile(worddoclines):
#    return createTextFile(worddoclines)


# def createTxtFile(worddoclines, filename="txtfile.txt"):
# txtfile=open('docs\\txt-output.txt','w')

def enc_E000createTextFile(worddoclines, filename="txtfile.txt"):
    txtfile = open(filename, 'w')

    for tup in worddoclines:
        ln = tup[0]  # ln=text
        colorcat = tup[1]  # colorcategory

        temp = ln[0:9]
        if temp == "linebreak":
            ln = "";

        if colorcat == 1:
            txtfile.write(ln + "\t****\n")
        else:
            txtfile.write(ln + "\n")

    txtfile.close()
    print("created txt06.txt now")

    return 1


# removed highlight options, keep for refernece
# highlightcolor1=WordHighlightColor.PINK, \
# highlightcolor2=WordHighlightColor.BRIGHT_GREEN,\
# highlightcolor3=WordHighlightColor.TURQUOISE):

# filename="docs\\word-file.docx"

def enc_D000createWordFile( \
        enc_B009worddoclines,
        enc_B010colorcatAuto=FontColor.BLACK, \
        enc_B001colorcat1=FontColor.RED, \
        enc_B002colorcat2=FontColor.GREEN, \
        enc_B003colorcat3=FontColor.BLUE,
        enc_B004colorcat4=FontColor.PURPLE,
        enc_B005colorcat5=FontColor.GOLD2,
        enc_B006colorcat6=FontColor.SILVER,
        enc_B007colorcat7=FontColor.ORANGE,
        enc_B008filename="docs\\word-file.docx"
):
    print("creating word doc now")
    enc_D021colsArray = (enc_B010colorcatAuto, enc_B001colorcat1, enc_B002colorcat2, enc_B003colorcat3, \
                         enc_B004colorcat4, enc_B005colorcat5, enc_B006colorcat6, enc_B007colorcat7)

    worddocument = Document()
    # worddocument.add_heading("Hello Word File",2)
    p = worddocument.add_paragraph("\n")
    p.add_run("\n\n").font_size = 8

    last_waslistitem = False

    for enc_D022tup in enc_B009worddoclines:
        ln = enc_D022tup[0] + "\n"
        # ln="\n"+tup[0] #no, problems
        if last_waslistitem:
            ln = "\n" + ln
            pass

        enc_D023colorcat = enc_D022tup[1]  # color category

        enc_D024size = int(enc_D022tup[2]) - 6  ## reduce size
        enc_D024bold = enc_D022tup[3]
        enc_D024italic = enc_D022tup[4]

        enc_D024msword_heading = enc_D022tup[5]
        enc_D024msword_listitem = enc_D022tup[6]
        enc_D024msword_highlight = enc_D022tup[7]

        if enc_D024size >= 10:
            enc_D024size = enc_D024size - 2

        # print("highlightcolor=",hc)
        # p.add_run("").font_size=size.value
        # obj_font.size = Pt(10)
        # p.font_size=Pt(size.value)

        # font_styles = worddocument.styles
        # font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        # font_object.size = Pt(size.value)

        # print(ln[:9], ln[:9]=="pagebreak" )

        # cammot set both bold and color, so bold takes precedence, fix later
        # cannot change fonzt size either

        last_waslistitem = False
        temp = ln[:9]
        if temp == "pagebreak":
            # print("new page here")
            worddocument.add_page_break()
            # p=worddocument.add_paragraph("NEW PAGE  007\n")
            p = worddocument.add_paragraph("\n")

            continue
            pass

        if temp == "linebreak":
            txt = p.add_run("\n")
            continue
            pass

        # print(ln,"highlight=",msword_highlight)
        if enc_D024msword_highlight:
            # txt.font.highlight_color = WordHighlightColor.YELLOW.value
            # print(ln, "highlight this")
            # processed below
            pass

        # validate in insert_line
        if enc_D024msword_heading >= 1:
            worddocument.add_heading(ln, enc_D024msword_heading)
            p = worddocument.add_paragraph("\n")

            continue
            pass

        if enc_D024msword_listitem:
            # ln="* "+ln # just add star in front , at moment

            ln = ln.strip()  # remove new line space,
            p = worddocument.add_paragraph(ln, style="List Bullet")

            # p.font.size=Pt(enc_D024size), # no

            last_waslistitem = True
            if enc_D024msword_highlight:
                # p.font.highlight_color = WordHighlightColor.YELLOW.value # causes error
                # p.highlight_color = WordHighlightColor.YELLOW.value # no error, but doesnt't work
                pass

            continue
            pass

        # **  WordFile  ***
        # colorChoice=FontColor.BLACK
        enc_D034colorChoice = enc_D021colsArray[enc_D023colorcat]

        enc_D034red = enc_D034colorChoice[0]
        enc_D034green = enc_D034colorChoice[1]
        enc_D034blue = enc_D034colorChoice[2]

        txt = p.add_run(ln)
        txt.bold = enc_D024bold
        txt.italic = enc_D024italic

        txt.font.size = Pt(enc_D024size)
        txt.font.color.rgb = RGBColor(enc_D034red, enc_D034green, enc_D034blue)

        if enc_D024msword_highlight == False:
            pass
        elif enc_D024msword_highlight == 1:
            txt.font.highlight_color = WordHighlightColor.YELLOW.value
        elif enc_D024msword_highlight == 2:
            txt.font.highlight_color = WordHighlightColor.BRIGHT_GREEN.value
        elif enc_D024msword_highlight == 3:
            txt.font.highlight_color = WordHighlightColor.TEAL.value
        elif enc_D024msword_highlight == 4:
            txt.font.highlight_color = WordHighlightColor.PINK.value

    print("saving word doc", enc_B008filename)
    # word_filename="csv\\docs\\worddoc-output.docx"
    # filename="docs\\worddoc-output.docx"

    try:
        worddocument.save(enc_B008filename)
    except:
        print("exception occurred saving word file")
        if inputyesno("exception, occurred. --> Close Word file. Try again ?", True):
            worddocument.save(enc_B008filename)

    return 1


def getDoc():
    return MarianoDocument()


class MarianoList(list):
    pass


class MarianoDocument:

    def __init__(self):
        self.lines = MarianoList()
        pass

    def get_lines(self):
        return self.lines

    def clear_lines(self):
        self.lines.clear()

    def pagebreak(self, txt=""):
        self.insert_line("pagebreak")
        return 1

    def linebreak(self, txt=""):
        self.insert_line("linebreak")
        return 1

    def word_heading(self, *items, size=2):
        self.insert_line(*items, msword_heading=size)
        return 1

    def word_listitem(self, *items):
        self.insert_line(*items, msword_listitem=True)
        return 1

    def insert_line(self, *items, sep=" ", print_console=True, \
                    color_category=0, font_size=FontSize.MEDIUM, \
                    font_bold=False, font_italic=False, msword_listitem=False, \
                    msword_heading=0, msword_highlight=0, misc=0):
        # print("insert_doc, items=",*items)
        doclines = self.lines
        enc_A000doc_insert(doclines, *items, enc_A003sep=sep, enc_A004print_console=print_console, \
                           enc_A005color_category=color_category, enc_A006font_size=font_size.value, \
                           enc_A007font_bold=font_bold, enc_A008font_italic=font_italic, \
                           enc_A009msword_heading=msword_heading, enc_A010msword_listitem=msword_listitem, \
                           enc_A011msword_highlight=msword_highlight, misce=misc)
        return 1

    def createPDFFile(self, filename="pdf-file.pdf", \
                      colorcatAuto=FontColor.BLACK, \
                      colorcat1=FontColor.RED, \
                      colorcat2=FontColor.GREEN, \
                      colorcat3=FontColor.BLUE,
                      colorcat4=FontColor.PURPLE,
                      colorcat5=FontColor.GOLD2,
                      colorcat6=FontColor.SILVER,
                      colorcat7=FontColor.ORANGE

                      ):
        doclines = self.lines
        return enc_B000createPDFFile(enc_B009worddoclines=doclines, enc_B010colorcatAuto=colorcatAuto,
                                     enc_B001colorcat1=colorcat1, enc_B002colorcat2=colorcat2,
                                     enc_B003colorcat3=colorcat3, enc_B004colorcat4=colorcat4,
                                     enc_B005colorcat5=colorcat5, enc_B006colorcat6=colorcat6,
                                     enc_B007colorcat7=colorcat7, enc_B008filename=filename)

        ## ends createPDF

    def createTextFile(self, filename="txtfile.txt"):
        doclines = self.lines
        return enc_E000createTextFile(doclines, filename)

    def createWordFile(self, filename="word-file.docx", \
                       colorcatAuto=FontColor.BLACK, \
                       colorcat1=FontColor.RED, \
                       colorcat2=FontColor.GREEN, \
                       colorcat3=FontColor.BLUE,
                       colorcat4=FontColor.PURPLE,
                       colorcat5=FontColor.GOLD2,
                       colorcat6=FontColor.SILVER,
                       colorcat7=FontColor.ORANGE

                       ):
        doclines = self.lines
        return enc_D000createWordFile(enc_B009worddoclines=doclines, enc_B010colorcatAuto=colorcatAuto,
                                      enc_B001colorcat1=colorcat1, enc_B002colorcat2=colorcat2,
                                      enc_B003colorcat3=colorcat3, enc_B004colorcat4=colorcat4,
                                      enc_B005colorcat5=colorcat5, enc_B006colorcat6=colorcat6,
                                      enc_B007colorcat7=colorcat7, enc_B008filename=filename)

    pass















